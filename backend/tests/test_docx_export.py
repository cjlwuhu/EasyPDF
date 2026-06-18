from io import BytesIO

from docx import Document as DocxDocument
from fastapi import HTTPException
import pytest

from app.api.routes.documents import get_translation_docx
from app.db.models import Document, Paragraph, ParagraphStatus
from app.services.docx_export import build_translation_docx, safe_docx_filename


def test_docx_contains_title_and_nonempty_translations():
    payload = build_translation_docx("Paper", ["第一段译文", "  ", "第二段译文"])

    document = DocxDocument(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Paper" in text
    assert "第一段译文" in text
    assert "第二段译文" in text
    assert len([paragraph for paragraph in document.paragraphs if paragraph.text]) == 3


def test_docx_removes_xml_incompatible_control_characters():
    payload = build_translation_docx("Paper", ["First\x00 paragraph", "Second\x0b paragraph"])

    document = DocxDocument(BytesIO(payload))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "First paragraph" in text
    assert "Second paragraph" in text


def test_safe_docx_filename_removes_windows_reserved_characters():
    assert safe_docx_filename('A/B:C*?') == "A_B_C___translated.docx"


def test_safe_docx_filename_falls_back_for_empty_title():
    assert safe_docx_filename(' ... ') == "translation_translated.docx"


def add_document(db_session, translated_text: str) -> Document:
    document = Document(
        title="Paper",
        original_filename="paper.pdf",
        file_path="paper.pdf",
        status="parsed",
    )
    document.paragraphs.append(
        Paragraph(
            page_number=1,
            order_index=0,
            source_text="English source",
            translated_text=translated_text,
            status=ParagraphStatus.translated if translated_text else ParagraphStatus.pending,
        )
    )
    db_session.add(document)
    db_session.commit()
    db_session.refresh(document)
    return document


def test_translation_docx_route_rejects_document_without_translations(db_session):
    document = add_document(db_session, "")

    with pytest.raises(HTTPException) as exc_info:
        get_translation_docx(document.id, db_session)

    assert exc_info.value.status_code == 409


def test_translation_docx_route_returns_download_response(db_session):
    document = add_document(db_session, "中文译文")

    response = get_translation_docx(document.id, db_session)

    assert response.media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert "Paper_translated.docx" in response.headers["content-disposition"]


def test_translation_docx_route_supports_unicode_document_title(db_session):
    document = add_document(db_session, "Translated text")
    document.title = "\u82f1\u8bed\u8bba\u6587"
    db_session.commit()

    response = get_translation_docx(document.id, db_session)
    content_disposition = response.headers["content-disposition"]

    content_disposition.encode("latin-1")
    assert 'filename="translation.docx"' in content_disposition
    assert "filename*=UTF-8''" in content_disposition
