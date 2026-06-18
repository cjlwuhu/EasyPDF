from io import BytesIO
import re

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


XML_INCOMPATIBLE_CHARACTERS = re.compile(
    r"[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]"
)


def _xml_safe_text(value: str) -> str:
    return XML_INCOMPATIBLE_CHARACTERS.sub("", value)


def build_translation_docx(title: str, translations: list[str]) -> bytes:
    title = _xml_safe_text(title)
    document = DocxDocument()
    document.core_properties.title = title

    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(11)
    normal_style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.paragraph_format.line_spacing = 1.55
    normal_style.paragraph_format.space_after = Pt(8)

    heading = document.add_heading(title, level=0)
    heading.paragraph_format.space_after = Pt(18)
    for run in heading.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for translation in translations:
        normalized = _xml_safe_text(translation).strip()
        if normalized:
            document.add_paragraph(normalized)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def safe_docx_filename(title: str) -> str:
    safe_title = re.sub(r'[<>:"/\\|?*]', "_", title).strip(" .") or "translation"
    return f"{safe_title}_translated.docx"
